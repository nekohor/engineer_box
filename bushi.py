# -*- coding: utf-8 -*-
from jinja2 import Template
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches


class Templet():

    def __init__(self, template_name):
        """
        initialization for read the template
        """
        # self.env = Environment(loader=PackageLoader(
        #     package, path))
        # self.template = self.env.get_template(template_name)
        with open(template_name, 'r', encoding='UTF-8') as f:
            self.template = Template(f.read())

    def put(self, doc, token, kwdict):
        # 分级标题
        for i in range(5):
            level = i + 1
            if token.startswith('<h%d>' % level):
                words = token[4:-5].strip()
                doc.add_heading(words, level)
        # 状态处理
        if token.startswith('<p>'):
            words = token[3:-4].strip()
            doc.add_paragraph(words)
        elif token.startswith('<span>'):
            words = token[6:-7].strip()
            doc.add_paragraph(words)
        elif token.startswith('<table>'):
            words = token[7:-8].strip()
            self.add_one_table(doc, kwdict[words])
        elif token.startswith('<picture>'):
            words = token[9:-10].strip()
            doc.add_picture(words, width=Inches(5.2))
        else:
            pass
            # print("Don't understand tag : %s" % token)

    def render(self, doc, kwdict):
        self.render_result = self.template.render(kwdict)
        # put to docx
        for token in self.render_result.split("\n"):
            # print(token)
            self.put(doc, token, kwdict)

    def add_one_table(self, document, summary):
        row_num = len(summary.index) + 1
        col_num = len(summary.columns) + 1
        print(row_num, col_num)
        table = document.add_table(rows=row_num, cols=col_num)
        # 表格第一行为DataFrame的列名
        hdr_cells = table.rows[0].cells
        side_cells = table.columns[0].cells
        i = 1
        for val in summary.columns:
            hdr_cells[i].text = str(val)
            i = i + 1

        i = 1
        for val in summary.index:
            side_cells[i].text = str(val)
            i = i + 1

        i = 1
        for col in summary.columns:
            j = 1
            for idx in summary.index:
                table.columns[i].cells[j].text = (
                    str(summary.loc[idx, col])
                    # 注意cell的text只能接收字符串
                )
                j = j + 1
            i = i + 1
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        document.add_paragraph("")
